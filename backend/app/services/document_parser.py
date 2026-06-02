"""Document parsing helpers for uploaded knowledge-base files."""
import csv
import io
from dataclasses import dataclass, field
from html import unescape
from typing import Any, Dict, List, Optional

from app.core.exceptions import AppException


@dataclass
class ParsedDocument:
    content: str
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    """Parse uploaded files into retrieval-friendly plain/Markdown text."""

    _TEXT_EXTENSIONS = {"txt", "md"}

    def parse(self, file_bytes: bytes, ext: str, filename: Optional[str] = None) -> ParsedDocument:
        ext = (ext or "").lower()

        if ext in self._TEXT_EXTENSIONS:
            return self._parse_text(file_bytes, ext)
        if ext == "pdf":
            return self._parse_pdf(file_bytes)
        if ext == "docx":
            return self._parse_docx(file_bytes)
        if ext == "csv":
            return self._parse_csv(file_bytes)
        if ext in {"xlsx", "xlsm"}:
            return self._parse_xlsx(file_bytes)
        if ext == "doc":
            return ParsedDocument(
                content="",
                metadata={"parser": "unsupported", "warning": "旧版 .doc 二进制格式暂不支持直接解析，请转换为 .docx 或 PDF。"},
            )

        return ParsedDocument(
            content="",
            metadata={"parser": "unsupported", "warning": f"暂不支持解析 {ext or 'unknown'} 格式。"},
        )

    def _parse_text(self, file_bytes: bytes, ext: str) -> ParsedDocument:
        text = self._decode_text(file_bytes)
        return ParsedDocument(content=text, metadata={"parser": ext})

    def _parse_pdf(self, file_bytes: bytes) -> ParsedDocument:
        try:
            import pdfplumber
        except ImportError:
            return self._parse_pdf_with_pypdf(file_bytes)

        page_texts = []
        table_count = 0
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                parts = []
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())

                for table_index, table in enumerate(page.extract_tables() or [], start=1):
                    markdown = self._table_to_markdown(table)
                    if markdown:
                        table_count += 1
                        parts.append(f"[Page {page_index} Table {table_index}]\n{markdown}")

                if parts:
                    page_texts.append(f"## Page {page_index}\n" + "\n\n".join(parts))

        return ParsedDocument(
            content="\n\n".join(page_texts),
            metadata={"parser": "pdfplumber", "pages": len(page_texts), "tables": table_count},
        )

    def _parse_pdf_with_pypdf(self, file_bytes: bytes) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise AppException("PDF解析依赖缺失，请安装 pdfplumber 或 pypdf", 500) from exc

        reader = PdfReader(io.BytesIO(file_bytes))
        page_texts = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                page_texts.append(f"## Page {page_index}\n{text.strip()}")

        return ParsedDocument(
            content="\n\n".join(page_texts),
            metadata={"parser": "pypdf", "pages": len(page_texts), "tables": 0},
        )

    def _parse_docx(self, file_bytes: bytes) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise AppException("DOCX解析依赖缺失，请安装 python-docx", 500) from exc

        document = DocxDocument(io.BytesIO(file_bytes))
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        table_count = 0
        for table_index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            markdown = self._table_to_markdown(rows)
            if markdown:
                table_count += 1
                parts.append(f"[Table {table_index}]\n{markdown}")

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"parser": "python-docx", "tables": table_count},
        )

    def _parse_csv(self, file_bytes: bytes) -> ParsedDocument:
        text = self._decode_text(file_bytes)
        rows = list(csv.reader(io.StringIO(text)))
        markdown = self._table_to_markdown(rows)
        return ParsedDocument(
            content=markdown or text,
            metadata={"parser": "csv", "rows": len(rows), "tables": 1 if markdown else 0},
        )

    def _parse_xlsx(self, file_bytes: bytes) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise AppException("Excel解析依赖缺失，请安装 openpyxl", 500) from exc

        workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts = []
        table_count = 0
        for sheet in workbook.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value).strip() for value in row]
                if any(values):
                    rows.append(values)
            markdown = self._table_to_markdown(rows)
            if markdown:
                table_count += 1
                parts.append(f"## Sheet: {sheet.title}\n{markdown}")

        return ParsedDocument(
            content="\n\n".join(parts),
            metadata={"parser": "openpyxl", "sheets": len(workbook.worksheets), "tables": table_count},
        )

    def _decode_text(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
            try:
                return file_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="ignore")

    def _table_to_markdown(self, rows: List[List[Any]]) -> str:
        clean_rows = [
            [self._clean_cell(cell) for cell in row]
            for row in rows
            if row and any(self._clean_cell(cell) for cell in row)
        ]
        if not clean_rows:
            return ""

        width = max(len(row) for row in clean_rows)
        normalized = [row + [""] * (width - len(row)) for row in clean_rows]
        header = normalized[0]
        body = normalized[1:] or [[""] * width]

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * width) + " |",
        ]
        lines.extend("| " + " | ".join(row) + " |" for row in body)
        return "\n".join(lines)

    def _clean_cell(self, value: Any) -> str:
        if value is None:
            return ""
        text = unescape(str(value)).replace("\r", " ").replace("\n", " ").strip()
        return text.replace("|", "\\|")
